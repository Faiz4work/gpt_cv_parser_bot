from docx import Document
import os

def read_doc(file_path):
    # file_path = "JSON.docx"
    # file_path = "Muhammad Anas CV.docx"
    if not os.path.exists(file_path):
        print("File not found!")
        return None

    doc = Document(file_path)
    document_text = [t.text for t in doc.paragraphs]
    document_text = " ".join(document_text)
    return document_text

    # for paragraph in doc.paragraphs:
    #     print(paragraph.text)

    # print(len(doc.paragraphs))





